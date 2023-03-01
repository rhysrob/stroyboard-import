from django.shortcuts import redirect, render
from django.contrib import messages
from docx import Document
import json
from .utils import (add_static_card_content, loop_though_table_content, loop_through_tables_in_cells,append_table_to_answer_hint_extension)

# Create your views here.
def scraper(request):
    files = request.FILES.getlist('filename')
    dic = {}
    lesson_number = 0
    screen_number = 0
    screen_total = 0
    card_number = 0

    if request.POST:
        if not files:
            messages.error(request, 'The input is empty, please add a file')
        else:
            document = Document(files[0])
            tables = document.tables
            # code
            dic['code'] = tables[0].rows[0].cells[1].text.lower().replace('.',
                                                                          '-').strip()
            # This project name
            if "|" in tables[0].rows[1].cells[1].text:
                dic['resourceTitle'] = {
                    "en": tables[0].rows[1].cells[1].text.split('|')[0].strip(),
                    "cy": tables[0].rows[1].cells[1].text.split('|')[1].strip()
                }
            else:
                dic['resourceTitle'] = {
                    "en": tables[0].rows[1].cells[1].text.strip(),
                    "cy": ''
                }
            # Subject name
            if "|" in tables[0].rows[2].cells[1].text:
                dic['subject'] = {
                    "en": tables[0].rows[2].cells[1].text.split('|')[0].strip(),
                    "cy": tables[0].rows[2].cells[1].text.split('|')[1].strip()
                }
            else:
                dic['subject'] = {
                    "en": tables[0].rows[2].cells[1].text.strip(),
                    "cy": ''
                }
            # show bar
            dic['showSidebar'] = True
            # lessons
            dic['lessons'] = []

           # loop through tables
            for table in tables:
                if 'Language'.lower() in table.rows[0].cells[0].text.lower():
                    if "Lesson".lower() in table.rows[1].cells[0].text.lower():
                        try:
                            dic['lessons'].append(
                                {
                                    "title": {
                                        "en": table.rows[1].cells[1].text.strip(),
                                        "cy": table.rows[1].cells[2].text.strip(),
                                    },
                                    "active": False,
                                    "heroImage": "https://resource.download.wjec.co.uk/vtc/2022-23/el22-23_2-6/images/shakespeare-book.png",
                                    "progress": 0,
                                    "subLessons": []
                                }
                            )

                            lesson_number += 1
                            screen_number = 0

                            if "Lesson".lower() in table.rows[1].cells[1].text.lower():
                                messages.error(
                                    request, f' Lesson {lesson_number} tile is not correct -> {table.rows[1].cells[1].text}')
                                continue
                        except IndexError:
                            messages.error(
                                request, f'LESSON: {lesson_number}, screen: {screen_number} need to be added')

                        # adding screens
                        try:
                            if "Screen".lower() in table.rows[2].cells[0].text.lower():
                                try:
                                    dic['lessons'][lesson_number - 1]['subLessons'].append(
                                        {
                                            "title": {
                                                "en": table.rows[2].cells[1].text,
                                                "cy": table.rows[2].cells[2].text,
                                            },
                                            "slug": "slug",
                                            "cards": []
                                        }

                                    )
                                    screen_number += 1
                                    screen_total += 1
                                    card_number = 0
                                    if table.rows[2].cells[1].text == '':
                                        messages.error(
                                            request, f'{lesson_number}, {screen_number} is empty,  {table.rows[2].cells[0].text}')
                                    continue
                                except IndexError:
                                    messages.error(
                                        request, f'Screen in lesson: {lesson_number}, screen: {screen_number} need to be added')
                            continue
                        except IndexError:
                            messages.error(
                                request, f'Screen in lesson: {lesson_number}, screen: {screen_number} need to be added')

                    elif "Screen".lower() in table.rows[1].cells[0].text.lower():
                        try:
                            dic['lessons'][lesson_number - 1]['subLessons'].append(
                                {
                                    "title": {
                                        "en": table.rows[1].cells[1].text,
                                        "cy": table.rows[1].cells[2].text,
                                    },
                                    "slug": "slug",
                                    "cards": []
                                },

                            )
                            screen_number += 1
                            screen_total += 1
                            card_number = 0
                            continue
                        except IndexError:
                            messages.error(
                                request, f'Screen title in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # Lesson Objectives
                if 'LESSON OBJECTIVES'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "read",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": '',
                                        "cy": '',
                                        "_editorEn": "visual"
                                    }
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": []
                            },
                        )
                        card_number + 1
                        add_static_card_content(
                            dic, lesson_number, screen_number, card_number, table)
                        loop_though_table_content(
                            3, dic, lesson_number, screen_number, card_number, table)
                        loop_through_tables_in_cells(2, dic,lesson_number,screen_number,card_number,table, "content")
                        continue
                    except IndexError:
                        messages.error(
                            request, f'LESSON OBJECTIVES card in lesson: {lesson_number}, screen: {screen_number} need to be added ')

                 # adding reading cards
                if "reading".lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "read",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": '',
                                        "cy": '',
                                        "_editorEn": "visual"
                                    }
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": [

                                ]
                            },
                        )
                        card_number + 1
                        add_static_card_content(
                            dic, lesson_number, screen_number, card_number, table)
                        loop_through_tables_in_cells(2, dic,lesson_number,screen_number,card_number,table, "content")
                        continue
                    except IndexError:
                        messages.error(
                            request, f' Reading card in Lesson {lesson_number}, Screen {screen_number} need to add content')

                # Hero image block
                if 'Hero image block'.lower() in table.rows[0].cells[0].text.lower():
                    messages.error(
                        request, f'Hero image  block in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # Image block'
                if 'Image block'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "look",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": "",
                                        "cy": "",
                                        "_editorEn": "code"
                                    }
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": []
                            },
                        )
                        card_number + 1
                        messages.success(
                            request, f'Image block in {lesson_number} lessons and {screen_total} screens need image proccess')
                        add_static_card_content(
                            dic, lesson_number, screen_number, card_number, table)
                        loop_through_tables_in_cells(2, dic,lesson_number,screen_number,card_number,table, "content")

                        continue
                    except IndexError:
                        messages.error(
                            request, f'Image block card in lesson: {lesson_number}, screen: {screen_number} need to be added ')

                # adding WRITING cards
                if "WRITING".lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "write",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": '',
                                        "cy": '',
                                        "_editorEn": "visual"
                                    }
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": []
                            },
                        )
                        card_number + 1
                        add_static_card_content(
                            dic, lesson_number, screen_number, card_number, table)
                        loop_through_tables_in_cells(2, dic,lesson_number,screen_number,card_number,table, "content")
                        continue
                    except IndexError:
                        messages.error(
                            request, f'WRITING Card in lesson: {lesson_number}, screen: {screen_number} need to add content')

                # adding Question and Answer cards
                if "Question and Answer".lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "questionAndAnswer",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "question": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "answer": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "feedback": {
                                        "en": "",
                                        "cy": ""
                                    }
                                }
                            }
                        )
                        card_number + 1
                        add_static_card_content(
                            dic, lesson_number, screen_number, card_number, table, 'question')
                        try:
                            if table.rows[3].cells[0].text:
                                for i in range(3, len(table.rows)):
                                    dic['lessons'][lesson_number - 1]['subLessons'][screen_number -
                                                                                    1]['cards'][card_number - 1]['data']['answer']['en'] += f'{table.rows[i].cells[0].text}'
                                    dic['lessons'][lesson_number - 1]['subLessons'][screen_number -
                                                                                    1]['cards'][card_number - 1]['data']['answer']['cy'] += f'{table.rows[i].cells[1].text}'
                            messages.success(
                                request, f'questionandAnswer Card in lesson: {lesson_number}, screen: {screen_number} added correctly')
                        except IndexError:
                            continue
                        continue
                    except IndexError:
                        messages.error(
                            request, f'WRITING Card in lesson: {lesson_number}, screen: {screen_number} need to add content')

                # Adding CASE STUDY cards
                if "case study".lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "caseStudy",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": '',
                                        "cy": '',
                                        "_editorEn": "visual"
                                    }
                                },
                            },
                        )
                        card_number + 1
                        add_static_card_content(
                            dic, lesson_number, screen_number, card_number, table)
                        loop_through_tables_in_cells(2, dic,lesson_number,screen_number,card_number,table, "content")

                        continue
                    except IndexError:
                        messages.error(
                            request, f'Case study card in lesson: {lesson_number}, screen: {screen_number} need to be added ')

                # Adding SPEAKING cards
                if "speaking".lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "speaking",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": '',
                                        "cy": '',
                                        "_editorEn": "visual"
                                    }
                                },
                            },
                        )
                        card_number + 1
                        add_static_card_content(
                            dic, lesson_number, screen_number, card_number, table)
                        loop_through_tables_in_cells(2, dic,lesson_number,screen_number,card_number,table, "content")
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Speaking card in lesson: {lesson_number}, screen: {screen_number} need to add content ')

                # Adding SUMMARY cards
                if "summary".lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "summary",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": '',
                                        "cy": '',
                                        "_editorEn": "visual"
                                    }
                                },
                            },
                        )
                        card_number + 1
                        add_static_card_content(
                            dic, lesson_number, screen_number, card_number, table)
                        loop_through_tables_in_cells(2, dic,lesson_number,screen_number,card_number,table, "content")

                        continue
                    except IndexError:
                        messages.error(
                            request, f'summary card in lesson: {lesson_number}, screen: {screen_number} need to add content ')

                # Adding MCQ cards
                if "MCQ".lower() in table.rows[0].cells[0].text.lower():
                    try:
                        print(table.rows[0].cells[0].text)
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "mcq",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": '',
                                        "cy": '',
                                        "_editorEn": "visual"
                                    },
                                    "question":{
                                        "en":table.rows[i].cells[0].text,
                                        "cy":table.rows[i].cells[1].text
                                    },
                                    "choices":[
                                        {
                                            "answer":{"en":"","cy":""},
                                            "correct":"true"
                                        }
                                    ]
                                },
                            },
                        )
                        card_number + 1
                        add_static_card_content(
                            dic, lesson_number, screen_number, card_number, table)
                        loop_through_tables_in_cells(2, dic,lesson_number,screen_number,card_number,table, "content")

                        continue
                    except IndexError:
                        messages.error(
                            request, f'MCQ card in lesson: {lesson_number}, screen: {screen_number} need to add conntent ')

              # Spot the mistake card
                if "Spot the mistake".lower() in table.rows[0].cells[0].text.lower():
                    try:
                        for i in range(3, len(table.rows)):

                            dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                                {
                                    "type": "spotMistake",
                                    "data": {
                                        "title": {
                                            "en": "",
                                            "cy": ""
                                        },
                                        "content": {
                                            "en": "<p>Spot the mistake in the sentence and then insert the correction.</p>",
                                            "cy": "<p>Ceisiwch ddod o hyd i’r camgymeriad yn y frawddeg yna ewch ati i’w chywiro.</p>\n"
                                        },
                                        "question": {
                                            "en": table.rows[i].cells[0].text,
                                            "cy": table.rows[i].cells[2].text,
                                        },
                                        "answer": {
                                            "en": table.rows[i].cells[1].text,
                                            "cy": table.rows[i].cells[1].text,
                                        },
                                        "mistake": {
                                            "en": "<!** MANUALLY ADDD THE MISTAKE ** >",
                                            "cy": ""
                                        },
                                        "feedback": {
                                            "en": "",
                                            "cy": ""
                                        }
                                    },
                                    "answer": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "hint": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "extension": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "downloadFiles": []
                                },
                            )
                            card_number + 1
                            messages.error(
                                request, f'You need to add the Mistake manually in "spot the mistake activity" located in  lesson number: {lesson_number}, and screen {screen_number}')
                            continue
                    except IndexError:
                        messages.error(
                            request, f'Spot the mistake card in lesson: {lesson_number}, screen: {screen_number} need to be added ')

                # --------------------------------------------------------------------------------------------------------------------------------  check again video card
                if 'Video block'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "look",
                                "data": {

                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": f"\n<div class=\"video-container\"><iframe src=\"https://www.youtube.com/embed/here?rel=0&start=0&end=\" allowfullscreen></iframe></div>",
                                        "cy": f"\n<div class=\"video-container\"><iframe src=\"https://www.youtube.com/embed/here?rel=0&start=0&end=\" allowfullscreen></iframe></div>",
                                        "_editorEn": "code"
                                    }
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": []
                            },
                        )
                        card_number + 1
                        messages.success(
                            request, f'Need to add video id in {lesson_number} lessons and {screen_total} screens')
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Video block card in lesson: {lesson_number}, screen: {screen_number} need to be addded ')

                # Sortable into Columns
                # ------------------------------------------------------------------------------------------------------ list later loop through the content
                if 'Sortable into Columns'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        for i in range(3, len(table.rows)):
                            dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                                {
                                    "type": "sortableColumns",
                                    "data": {
                                        "title": {
                                            "en": '',
                                            "cy": '',
                                        },
                                        "content": {
                                            "en": "",
                                            "cy": ""
                                        },
                                        "dragItems": [
                                        ],
                                        "dropItems": [

                                        ]
                                    },
                                    "answer": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "hint": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "extension": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "downloadFiles": []
                                },

                            )
                            card_number + 1
                            add_static_card_content(
                                dic, lesson_number, screen_number, card_number, table)

                            # for dropItems
                            try:
                                for table in table.rows[3].cells[0].tables:
                                    for i in range(0, len(table.columns)):
                                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number - 1]['data']['dropItems'].append(
                                            {
                                                "id": f'col-{i}',
                                                "title": {
                                                    "en": table.columns[i].cells[0].text,
                                                    "cy": ''
                                                }
                                            }
                                        )
                                        # for dragItems
                                        for j in range(1, len(table.rows)):
                                            dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number - 1]['data']['dragItems'].append(
                                                {
                                                    "id": table.rows[j].cells[i].text,
                                                    "text": {
                                                        "en": table.rows[j].cells[i].text,
                                                        "cy": ''
                                                    },
                                                    "answer": f"col-{i}"
                                                },
                                            )
                            except IndexError:
                                continue
                            # For welsh
                            try:
                                for table in table.rows[3].cells[1].tables:
                                    for i in range(0, len(table.columns)):
                                        print('i value ->', table)
                                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number - 1]['data']['dropItems'][j]['title']['cy'] += table.columns[i].cells[0].text
                            except IndexError:
                                continue
                        messages.success(
                            request, f'Sortable into Columns card in lesson: {lesson_number}, screen: {screen_number} loaded successfully')
                    except IndexError:
                        messages.error(
                            request, f'Sortable into Columns card in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # 'REFLECTION card'
                if 'REFLECTION'.lower() in table.rows[0].cells[0].text:
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "reflection",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": '',
                                        "cy": '',
                                    }
                                }
                            }

                        )
                        card_number + 1
                        add_static_card_content(
                            dic, lesson_number, screen_number, card_number, table)
                        loop_through_tables_in_cells(2, dic,lesson_number,screen_number,card_number,table, "content")

                        continue
                    except IndexError:
                        messages.error(
                            request, f'REFLECTION card in Lesson {lesson_number}, Screen {screen_number}, need to be added')

              # --------------------------------------------------------------------------------------------------------------------need to be check Question carousel card
              #   need to loop for the tables
                if 'Question Carousel'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "carousel",
                                "data": {
                                        "title": {
                                            "en": "",
                                            "cy": ""
                                        },
                                    "content": {
                                            "en": "<p><span style=\"color:rgba(0,0,0,0.847);\">Write each question in your book, answer them and then share the answers with your teacher.&nbsp;</span></p>",
                                            "cy": "",
                                            "_editorEn": "visual"
                                            },
                                    "slides": [],
                                    "indicator": True,
                                    "loop": True
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": []
                            }

                        )
                        card_number + 1
                        try:
                            for i in range(3, len(table.rows)):
                                dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number - 1]['data']['slides'].append(
                                    {
                                        "en": f"<p><span style=\"color:rgba(0,0,0,0.847);\"><strong>{table.rows[i].cells[0].text}&nbsp;</strong></span></p>",
                                        "cy": f"<p><span style=\"color:rgba(0,0,0,0.847);\"><strong>{table.rows[i].cells[1].text}&nbsp;</strong></span></p>",
                                        "_editorEn": "visual"
                                    },
                                )
                            continue
                        except IndexError:
                            continue
                        continue
                    except IndexError:
                        messages.error(
                            request, f' Question carousel card in lesson: {lesson_number}, screen: {screen_number} need to be added ')

                 # OPINION
                if 'OPINION'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "opinionOpen",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": '',
                                        "cy": '',
                                    }
                                }
                            }
                        )
                        card_number + 1
                        add_static_card_content(
                            dic, lesson_number, screen_number, card_number, table)
                        loop_through_tables_in_cells(2, dic,lesson_number,screen_number,card_number,table, "content")
                        continue
                    except IndexError:
                        messages.error(
                            request, f'OPINION card in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # Multi-choice
                if 'Multi-choice with 1 correct answer'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "mcq",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": "<p>Select the answer you think is correct.</p>",
                                        "cy": "<p>Dewiswch bob ateb sy’n gywir, yn eich barn chi.</p>\n"
                                    },
                                    "question": {
                                        "en": table.rows[2].cells[0].text,
                                        "cy": table.rows[2].cells[1].text,
                                        "_editorEn": "visual"
                                    },
                                    "feedback": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "choices": [
                                    ],
                                    "disableShuffleStatements": False
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": []
                            },

                        )
                        card_number + 1
                        try:
                            for i in range(3, len(table.rows)):
                                if table.rows[3].cells[0].text == table.rows[i].cells[0].text:
                                    dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number - 1]['data']['choices'].append(
                                        {
                                            "answer": {
                                                "en": table.rows[i].cells[0].text,
                                                "cy": table.rows[i].cells[1].text
                                            },
                                            "correct": True
                                        }
                                    )
                                else:
                                    dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number - 1]['data']['choices'].append(
                                        {
                                            "answer": {
                                                "en": table.rows[i].cells[0].text,
                                                "cy": table.rows[i].cells[1].text
                                            },
                                            "correct": False
                                        }
                                    )
                            continue
                        except IndexError:
                            continue
                    except IndexError:
                        messages.error(
                            request, f'Multi-choice with 1 correct answer card in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # ------------------------------------------------------------------------------------------------------Done loop through the content
                # Fill the gaps drodown cards
                if 'Fill the gaps – dropdown'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "fillBlankDropdown",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "activityContent": {
                                        "en": '',
                                        "cy": "",
                                        "_editorEn": "code"
                                    },
                                    "options": {
                                        "en": [],
                                        "cy": []
                                    }
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": []
                            }

                        )
                        card_number + 1
                        try:
                            for i in range(2, len(table.rows)):
                                dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number -
                                                                                                            1]['data']['activityContent']['en'] += f'<p>{table.rows[i].cells[0].text}</p>\n'
                                dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number -
                                                                                                            1]['data']['activityContent']['cy'] += f'<p>{table.rows[i].cells[2].text}</p>\n'
                                # appending the answers
                                dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number -
                                                                                                            1]['data']['options']['en'].append(table.rows[i].cells[1].text.split("\u00a0"))
                                dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number -
                                                                                                            1]['data']['options']['cy'].append(table.rows[i].cells[1].text.split("\u00a0"))
                            messages.error(
                                request, f'Fill the gaps – dropdown card in lesson: {lesson_number}, screen: {screen_number} need to add dropdown')
                        except IndexError:
                            continue
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Fill the gaps – dropdown card in lesson: {lesson_number}, screen: {screen_number} need to be added')

                 # Fill the gaps typing card
                # ------------------------------------------------------------------------------------------------------ loop through the content
                if 'Fill the gaps - Typing'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "fillBlankText",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "activityContent": {
                                        "en": "",
                                        "cy": "",
                                        "_editorEn": "code"
                                    }
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": []
                            },
                        )
                        card_number + 1
                        try:
                            for i in range(2, len(table.rows)):
                                dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number -
                                                                                                            1]['data']['activityContent']['en'] += f'<p>{table.rows[i].cells[0].text}</p>\n'
                                dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number -
                                                                                                            1]['data']['activityContent']['cy'] += f'<p>{table.rows[i].cells[1].text}</p>\n'
                            messages.error(
                                request, f'Fill the gaps – Typing card in lesson: {lesson_number}, screen: {screen_number} need to add input')
                        except IndexError:
                            continue
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Fill the gaps – Typing card in lesson: {lesson_number}, screen: {screen_number} need to be added')

                 # Question & Answer card
                if 'Question & Answer'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "questionAndAnswer",
                                        "data": {
                                            "title": {
                                                "en": "",
                                                "cy": ""
                                            },
                                            "content": {
                                                "en": "Additional content",
                                                "cy": "CY_Aditional content"
                                            },
                                            "question": {
                                                "en": table.rows[2].cells[0].text,
                                                "cy": table.rows[2].cells[1].text,
                                            },
                                            "answer": {
                                                "en": table.rows[3].cells[0].text,
                                                "cy": table.rows[3].cells[1].text,
                                            },
                                            "feedback": {
                                                "en": "Feedback",
                                                "cy": "CY_Feedback"
                                            }
                                        }
                            }
                        )
                        card_number + 1
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Question & Answer card in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # Random question generator with answer
                # ------------------------------------------------------------------------------------------------------ loop through the content
                if 'Random question generator with answer'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "randomGenerator",
                                "data": {
                                    "title": {
                                        "en": "Random Question Generator - Answers",
                                        "cy": "Random Question Generator - Answers CY"
                                    },
                                    "showAnswers": True,
                                    "questions": [
                                        {
                                            "question": {
                                                "en": table.rows[2].cells[0].text,
                                                "cy": table.rows[2].cells[1].text,
                                            },
                                            "answer": {
                                                "en": table.rows[3].cells[0].text,
                                                "cy": table.rows[3].cells[1].text,
                                            }
                                        },
                                    ]
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": []
                            }
                        )
                        card_number + 1
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Random question generator with answer card in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # Random question generator -no answer
                # ------------------------------------------------------------------------------------------------------ loop through the content
                if 'Random question generator - no answer'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "randomGenerator",
                                "data": {
                                    "title": {
                                        "en": "Random Question Generator - Answers",
                                        "cy": "Random Question Generator - Answers CY"
                                    },
                                    "content": {
                                        "en": f"<p>{table.rows[2].cells[0].text}</p>",
                                        "cy": f"<p>{table.rows[2].cells[1].text}</p>",
                                    },
                                    "showAnswers": True,
                                    "questions": [
                                        {
                                            "question": {
                                                "en": "<p>Question one</p>",
                                                "cy": "<p>Cwestiwn un</p>"
                                            },
                                            "answer": {
                                                "en": "Answer one",
                                                "cy": "Ateb un"
                                            }
                                        },
                                        {
                                            "question": {
                                                "en": "<p>Question two</p>",
                                                "cy": "<p>Cwestiwn dau</p>"
                                            },
                                            "answer": {
                                                "en": "Answer two",
                                                "cy": "Ateb dau"
                                            }
                                        },
                                        {
                                            "question": {
                                                "en": "<p>Question three</p>",
                                                "cy": "<p>Cwestiwn tri</p>"
                                            },
                                            "answer": {
                                                "en": "Answer three",
                                                "cy": "Ateb tri"
                                            }
                                        }
                                    ]
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": []
                            }
                        )
                        card_number + 1
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Random question generator - no answer card in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # Image Carousel block
                if 'Image Carousel block'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "carousel",
                                "data": {
                                    "scrollBar": False,
                                    "indicator": True,
                                    "loop": True,
                                    "slides": [
                                        {
                                            "en": f"<p>{table.rows[2].cells[0].text}</p>",
                                            "cy": f"<p>{table.rows[2].cells[1].text}</p>",
                                        }
                                    ]
                                }
                            },
                        )
                        card_number + 1
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Image Carousel block card in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # Sound block
                if 'Sound block'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "sound",
                                "data": {
                                    "title": {
                                        "en": "Sound/Audio",
                                        "cy": "_Sound/Audio"
                                    },
                                    "content": {
                                        "en": f"<a href='{table.rows[2].cells[0].text}' target='_blank'><img src='https://resource.download.wjec.co.uk/vtc/2020-21/bl20-21_4-26/bbc-audio.jpg' class='img-block'></a>",
                                        "cy": f"<a href='{table.rows[2].cells[1].text} target='_blank'><img src='https://resource.download.wjec.co.uk/vtc/2020-21/bl20-21_4-26/bbc-audio.jpg' class='img-block'></a>"
                                    }
                                }
                            }
                        )
                        card_number + 1
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Sound block in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # Download block
                # ------------------------------------------------------------------------------------------------------ Json errror
                if 'Download block'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number - 1]['downloadFiles'].append(
                            {
                                "fileName": {
                                    "en": table.rows[2].cells[0].text,
                                    "cy": table.rows[2].cells[1].text
                                },
                                "url": {
                                    "en": 'Put download link here',
                                    "cy": 'Cy Put download link here',
                                }
                            }
                        )
                        card_number + 1
                        messages.success(
                            request, f'Download block in lesson: {lesson_number}, screen: {screen_number} loaded successfully')
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Download block card in lesson: {lesson_number}, screen: {screen_number} need to be added')
                    except KeyError:
                        messages.error(
                            request, f'Download block card in lesson: {lesson_number}, screen: {screen_number} errors')

                # HINT/SUPPORT
                # ------------------------------------------------------------------------------------------------------ Json errror
                if 'HINT/SUPPORT'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number -
                                                                        1]['cards'][card_number - 1]['hint']['en'] += f'<p>{table.rows[2].cells[0].text}</p>'
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number -
                                                                        1]['cards'][card_number - 1]['hint']['cy'] += f'<p>{table.rows[2].cells[1].text}</p>'
                        append_table_to_answer_hint_extension(2, dic,lesson_number,screen_number,card_number, table, "hint")                            
                        card_number + 1
                        messages.success(
                            request, f'HINT/SUPPORT in lesson: {lesson_number}, screen: {screen_number} loaded successfully')
                        continue
                    except IndexError:
                        messages.error(
                            request, f'HINT/SUPPORT card in lesson: {lesson_number}, screen: {screen_number} need to be added')
                    except KeyError:
                        messages.error(
                            request, f'HINT/SUPPORT card in lesson: {lesson_number}, screen: {screen_number} errors')

                # GLOSSARY
                # ------------------------------------------------------------------------------------------------------ loop in welsh through the content
                if 'GLOSSARY'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        for i in range(2, len(table.rows)):
                            if table.rows[i].cells[0].text.lower().strip() in dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number - 1]['data']['content']['en'].lower():
                                dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number -
                                                                                                            1]['data']['content']['en'] += f'<span data-glossary=\"{table.rows[i].cells[1].text}\">{table.rows[i].cells[0].text}</span>\n'
                            if table.rows[i].cells[2].text.lower().strip() in dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number - 1]['data']['content']['cy'].lower():
                                dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number -
                                                                                                            1]['data']['content']['cy'] += f'<span data-glossary=\"{table.rows[i].cells[3].text}\">{table.rows[i].cells[2].text}</span>\n'
                            messages.success(
                                request, f'GLOSSARY card in lesson: {lesson_number}, screen: {screen_number} loaded correctly')
                    except IndexError:
                        messages.error(
                            request, f'GLOSSARY {lesson_number}, screen: {screen_number} need to be in the correct position')
                        continue

                # EXTENSION
                # ------------------------------------------------------------------------------------------------------ Json errror
                if 'EXTENSION'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number -
                                                                                                    1]['extension']['en'] += f'<p>{table.rows[2].cells[0].text}</p>'
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number -
                                                                                                    1]['extension']['cy'] += f'<p>{table.rows[2].cells[1].text}</p>'
                                                                                                    
                        card_number + 1
                        append_table_to_answer_hint_extension(2, dic,lesson_number,screen_number,card_number, table, "extension")                            

                        messages.success(
                            request, f'EXTENSION in lesson: {lesson_number}, screen: {screen_number} loaded successfully')
                        continue
                    except IndexError:
                        messages.error(
                            request, f'EXTENSION card in lesson: {lesson_number}, screen: {screen_number} need to be added')
                    except KeyError:
                        messages.error(
                            request, f'EXTENSION card in lesson: {lesson_number}, screen: {screen_number} errors')

                # Static text – ANSWER
                # ------------------------------------------------------------------------------------------------------ Json errror
                if 'Static text – ANSWER'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number -
                                                                        1]['cards'][card_number - 1]['answer']['en'] += f'<p>{table.rows[2].cells[0].text}</p>'
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number -
                                                                        1]['cards'][card_number - 1]['answer']['cy'] += f'<p>{table.rows[2].cells[1].text}</p>'
                        card_number + 1
                        append_table_to_answer_hint_extension(2, dic,lesson_number,screen_number,card_number, table, "answer")                            
                        messages.success(
                            request, f'ANSWER in lesson: {lesson_number}, screen: {screen_number} loaded successfully')
                        continue
                    except IndexError:
                        messages.error(
                            request, f'ANSWER card in lesson: {lesson_number}, screen: {screen_number} need to be added')
                    except KeyError:
                        messages.error(
                            request, f'ANSWER card in lesson: {lesson_number}, screen: {screen_number} errors')

                # TEST YOURSELF
                if 'TEST YOURSELF'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(

                            {
                                "type": "testYourself",
                                        "data": {
                                            "title": {
                                                "en": "",
                                                "cy": ""
                                            },
                                            "content": {
                                                "en": table.rows[2].cells[0].text,
                                                "cy": table.rows[2].cells[1].text
                                            },
                                        },
                                "downloadFiles": [

                                        ]
                            }
                        )
                        card_number + 1
                        continue
                    except IndexError or TypeError:
                        messages.error(
                            request, f'Static text – TEST YOURSELF (for past paper questions only) card in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # MARK YOURSELF
                if 'MARK YOURSELF'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(

                            {
                                "type": "markYourself",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": table.rows[2].cells[0].text,
                                        "cy": table.rows[2].cells[1].text
                                    },
                                },
                                "downloadFiles": [

                                ]
                            }

                        )
                        card_number + 1
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Mark Yourself Card": {lesson_number}, screen: {screen_number} need to be added')

                 # True or False
                # ------------------------------------------------------------------------------------------------------ list loop through the content
                if 'True or False'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "trueFalse",
                                "data": {
                                    "title": {
                                        "en": table.rows[1].cells[0].text,
                                        "cy": table.rows[1].cells[1].text
                                    },
                                    "content": {
                                        "en": table.rows[3].cells[0].text,
                                        "cy": table.rows[3].cells[1].text
                                    },
                                    "activityContent": [
                                        [
                                            {
                                                "str": {
                                                    "en": "A1",
                                                    "cy": "A1_CY"
                                                },
                                                "correct": True
                                            },
                                            {
                                                "str": {
                                                    "en": "A2",
                                                    "cy": "A2_CY"
                                                }
                                            }
                                        ],
                                        [
                                            {
                                                "str": {
                                                    "en": "B1",
                                                    "cy": "B1_Cy"
                                                },
                                                "correct": True
                                            },
                                            {
                                                "str": {
                                                    "en": "B2",
                                                    "cy": "B2_CY"
                                                }
                                            }
                                        ],
                                        [
                                            {
                                                "str": {
                                                    "en": "C1",
                                                    "cy": "C1_CY"
                                                },
                                                "correct": True
                                            },
                                            {
                                                "str": {
                                                    "en": "C2",
                                                    "cy": "C2_CY"
                                                }
                                            }
                                        ],
                                        [
                                            {
                                                "str": {
                                                    "en": "D1",
                                                    "cy": "D1_Cy"
                                                },
                                                "correct": True
                                            },
                                            {
                                                "str": {
                                                    "en": "D2",
                                                    "cy": "D2_CY"
                                                }
                                            }
                                        ]
                                    ]
                                }
                            }
                        )
                        card_number + 1
                        continue
                    except IndexError:
                        messages.error(
                            request, f'True or False in lesson: {lesson_number}, screen: {screen_number} need to be added')

                 # Ranking correct answer
                 # ------------------------------------------------------------------------------------------------------ check this activity
                if 'Ranking correct answer'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "sortable",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": "<p>Reorder the items in to the correct order.</p>",
                                        "cy": "<p>CY Reorder the items in to the correct order.</p>"
                                    },
                                    "doCheck": True,
                                    "activityContent": [

                                    ]
                                }
                            }
                        )
                        card_number + 1
                        try:
                            if table.rows[2].cells[0].text:
                                for i in range(2, len(table.rows)):
                                    dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number - 1]['data']['activityContent'].append(
                                        {
                                            "str": {
                                                "en": table.rows[i].cells[0].text,
                                                "cy": table.rows[i].cells[1].text,
                                            },
                                        }
                                    )
                            messages.success(
                                request, f'Ranking correct answer card in lesson: {lesson_number}, screen: {screen_number} loaded correctly')

                        except IndexError:
                            continue
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Ranking correct card in lesson: {lesson_number}, screen: {screen_number} need to be added')


                 # Ranking no correct answer
                 # ------------------------------------------------------------------------------------------------------ check this activity
                if 'Ranking no'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "sortable",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": "<p>Amend the items in which you believe is correct.</p>",
                                        "cy": "<p>CY Amend the items in which you believe is correct.</p>"
                                    },
                                    "doCheck": False,
                                    "activityContent": [

                                    ]
                                }
                            }
                        )
                        card_number + 1
                        try:
                            if table.rows[2].cells[0].text:
                                for i in range(2, len(table.rows)):
                                    dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number - 1]['data']['activityContent'].append(
                                        {
                                            "str": {
                                                "en": table.rows[i].cells[0].text,
                                                "cy": table.rows[i].cells[1].text,
                                            },
                                        }
                                    )
                            messages.success(
                                request, f'Ranking no correct answer card in lesson: {lesson_number}, screen: {screen_number} loaded correctly')

                        except IndexError:
                            continue
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Ranking no correct card in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # Thought shower
                if 'Thought shower'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "thoughtShower",
                                "data": {
                                    "title": {
                                        "en": "",
                                        "cy": ""
                                    },
                                    "content": {
                                        "en": '',
                                        "cy": ''
                                    },
                                    "question": {
                                        "en": '',
                                        "cy": ''
                                    }
                                },
                                "answer": {
                                    "en": "",
                                    "cy": ""
                                },
                                "hint": {
                                    "en": "",
                                    "cy": ""
                                },
                                "extension": {
                                    "en": "",
                                    "cy": ""
                                },
                                "downloadFiles": []
                            }
                        )
                        card_number + 1
                        # looping and appending cells
                        loop_though_table_content(
                            3, dic, lesson_number, screen_number, card_number, table, 'question')
                        loop_through_tables_in_cells(2, dic,lesson_number,screen_number,card_number,table, "content")                               
                        # static appending
                        try:
                            dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'][card_number -
                                                                                                        1]['data']['content']['en'] += table.rows[2].cells[0].text
                            dic['lessons'][lesson_number - 1]['subLessons'][screen_number -
                                                                            1]['cards'][card_number - 1]['data']['content']['cy'] += table.rows[2].cells[1].text
                        except IndexError:
                            continue
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Thought shower in lesson: {lesson_number}, screen: {screen_number} need to be added')

                # Video with Question
                if 'Video with Question'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "videoQuestion",
                                "downloadFiles": [],
                                "hint": {},
                                "extension": {},
                                "answer": {},
                                "data": {
                                    "title": {
                                        "en": "Video with Question - Source",
                                        "cy": "_Read Title"
                                    },
                                    "content": {
                                        "en": "<p>Take a look at the following video and asnwer the questions that appear.</p>",
                                        "cy": "<p>Cymerwch olwg ar y fideo ac yna ateb y cewsitynau sy'n ymddangos.</p>"
                                    },
                                    "isYouTube": False,
                                    "path": f"{table.rows[2].cells[0].text}",
                                    "questions": [
                                        {
                                            "time": 3,
                                            "comment": {
                                                "en": "What do you think of this statement?",
                                                "cy": "Beth ydych yn meddwl o'r hyn?"
                                            }
                                        },
                                        {
                                            "time": 6,
                                            "comment": {
                                                "en": "Can you think of any other like this?",
                                                "cy": ""
                                            }
                                        },
                                        {
                                            "time": 9,
                                            "comment": {
                                                "en": "What do you think is missing here?",
                                                "cy": ""
                                            }
                                        }
                                    ]
                                }
                            },
                        )
                        card_number + 1
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Video with Question card with answer in Lesson {lesson_number}, Screen {screen_number}, add content {table.rows[0].cells[0].text}')

                # Gallery Card"
                # ------------------------------------------------------------------------------------------------------ list check this activity
                if 'Gallery Card'.lower() in table.rows[0].cells[0].text.lower():
                    try:
                        dic['lessons'][lesson_number - 1]['subLessons'][screen_number - 1]['cards'].append(
                            {
                                "type": "gallery",
                                "data": {
                                    "title": {
                                        "en": "Gallery",
                                        "cy": "Galeri"
                                    },
                                    "content": {
                                        "en": table.rows[2].cells[0].text,
                                        "cy": table.rows[2].cells[1].text
                                    },
                                    "photos": [
                                        {
                                            "title": {
                                                "en": "Image 1 title",
                                                "cy": "_Image 1 title"
                                            },
                                            "content": {
                                                "en": "Caption 1",
                                                "cy": "_Caption"
                                            },
                                            "source": "https://resource.download.wjec.co.uk/vtc/2020-21/bL20-21_1-1/allegorical-portrait.jpg"
                                        },
                                        {
                                            "title": {
                                                "en": "Image 2 title",
                                                "cy": "_Image 2 title"
                                            },
                                            "content": {
                                                "en": "Caption 2",
                                                "cy": "_Caption"
                                            },
                                            "source": "https://resource.download.wjec.co.uk/vtc/2020-21/bL20-21_1-1/armada-portrait.jpg"
                                        },
                                    ]
                                }
                            }
                        )
                        card_number + 1
                        continue
                    except IndexError:
                        messages.error(
                            request, f'Gallery Card in lesson: {lesson_number}, screen: {screen_number} need to be added')

            messages.success(
                request, f'File converted succesfully, it has {lesson_number} lessons and {screen_total} screens')
    dic = json.dumps(dic, indent=4)
    context = {'dic': dic}
    return render(request, 'scraper.html', context)
